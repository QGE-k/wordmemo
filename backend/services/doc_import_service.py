"""
文档导入解析服务
支持 txt / docx / xlsx / pdf 格式，提取单词并过滤非单词内容

核心目标：让多词短语、固定搭配、句型、谚语能够【完整】导入，
而不是被拆分成零散的单词（如 "be accustomed to doing sth." 不能再被
拆成 "be" / "accustomed to doing sth."）。
"""
import re
import io

# 单词正则：允许字母、连字符、空格（复合词如 "sports meeting"）
# 不允许纯数字、特殊符号、中文
WORD_PATTERN = re.compile(
    r'^[a-zA-Z][a-zA-Z\-]*'           # 第一个字符必须是字母
    r'(?:\s[a-zA-Z][a-zA-Z\-]*)*'      # 后续可以是空格+字母（复合词）
    r'$'
)

# 短语/条目标题正则：允许字母、数字、空格、连字符、撇号（直/弯）、缩写点（如 sth.）、
# 逗号、问号、感叹号、冒号分号、省略号、括号、加号（如 "How are you?"、"no sooner...than..."、
# "out of sight, out of mind"、"sb1 charge sb2 money for sth"）
ENTRY_PATTERN = re.compile(r"^[a-zA-Z][a-zA-Z0-9\s\-'’\":.,?!…;()（）+]*$")

# 用来从任意文本中抽取候选 token（兜底用）
TOKEN_PATTERN = re.compile(r'[a-zA-Z][a-zA-Z\- ]{0,39}[a-zA-Z]')

# 常见无意义词/误识别噪声（仅用于过滤“单独成词”的填充词，
# 不影响多词短语的完整性）
NOISE_WORDS = {
    'the', 'a', 'an', 'and', 'or', 'but', 'of', 'to', 'in', 'on', 'at',
    'is', 'are', 'was', 'were', 'be', 'been', 'being', 'have', 'has', 'had',
    'do', 'does', 'did', 'will', 'would', 'can', 'could', 'should', 'may',
    'might', 'must', 'this', 'that', 'these', 'those', 'it', 'its', 'as',
    'for', 'with', 'by', 'from', 'about', 'into', 'through', 'during',
    'before', 'after', 'above', 'below', 'up', 'down', 'out', 'off',
    'over', 'under', 'again', 'further', 'then', 'once', 'here', 'there',
    'when', 'where', 'why', 'how', 'all', 'each', 'few', 'more', 'most',
    'other', 'some', 'such', 'no', 'nor', 'not', 'only', 'own', 'same',
    'so', 'than', 'too', 'very', 'just', 'also', 'page', 'chapter',
    'http', 'https', 'www', 'com', 'org', 'net', 'email', 'tel',
}

# 词性标记（如 "abandon v." -> 剥离 "v."，保留 "abandon"）
POS_MARKER_RE = re.compile(
    r'\s+(?:n|v|vt|vi|adj|adv|ad|prep|conj|pron|num|art|int|interj|aux|'
    r'modal|det|ger|part|colloc|phr|phrase)\.?\s*$',
    re.IGNORECASE
)

# 短语最大长度 / 短语允许的最大单词数（避免把整句当短语）
MAX_PHRASE_LEN = 60
MAX_PHRASE_WORDS = 10


def is_valid_word(text):
    """
    判断是否为有效英文单词
    - 长度 1-40
    - 字母开头
    - 允许连字符和空格（复合词）
    - 过滤噪声词
    - 复合词最多 3 个部分（2 个空格），避免整句被当成单词
    """
    if not text:
        return False
    text = text.strip()
    if len(text) < 1 or len(text) > 40:
        return False
    if not WORD_PATTERN.match(text):
        return False
    # 复合词空格数不能超过 2（最多 3 个部分），避免整句被当成单词
    if text.count(' ') > 2:
        return False
    # 全小写后检查噪声词（整体）
    if text.lower() in NOISE_WORDS:
        return False
    # 复合词的每个部分都不能是噪声词
    parts = text.split()
    for part in parts:
        if part.lower() in NOISE_WORDS:
            return False
    # 至少包含一个元音（避免类似 "bcdfg" 这种无意义字母组合）
    if not re.search(r'[aeiouAEIOU]', text):
        # 但允许少数无元音的合法词（如 my, by, fly, cry, rhythm）
        allowed_no_vowel = {'my', 'by', 'fly', 'cry', 'dry', 'gym', 'rhythm',
                            'why', 'shy', 'sky', 'sly', 'try', 'pty', 'nth'}
        if text.lower() not in allowed_no_vowel:
            return False
    return True


def is_valid_phrase(text):
    """
    判断是否为有效的多词短语（如 "sports meeting"、"in the long run"、
    "be accustomed to doing sth"、"It is high time that sb did sth"）
    - 2 个及以上单词
    - 只含字母、空格、连字符、撇号、缩写点
    - 不能全部是噪声词
    - 至少一个词包含元音
    """
    if not text:
        return False
    text = text.strip()
    if len(text) < 2 or len(text) > MAX_PHRASE_LEN:
        return False
    if not ENTRY_PATTERN.match(text):
        return False
    parts = [p for p in text.split() if p]
    if len(parts) < 2 or len(parts) > MAX_PHRASE_WORDS:
        return False
    # 不能全部是噪声词
    content = [p for p in parts if p.lower() not in NOISE_WORDS]
    if not content:
        return False
    # 至少一个词包含元音
    if not re.search(r'[aeiouAEIOU]', text):
        return False
    return True


def _clean_entry(english_part):
    """清洗英文部分：剥离词性标记和尾部残缺符号，返回规范条目标题"""
    s = english_part.strip()
    if not s:
        return ''
    # 剥离尾部词性标记（如 "abandon v." -> "abandon"）
    s = POS_MARKER_RE.sub('', s)
    # 剥离尾部残缺/孤立符号：空格、点、省略号、全角逗号顿号分号冒号、括号、加号
    # 保留句子尾部标点如 ? !（How are you?）以及内部逗号、撇号
    s = re.sub(r'[\s\.…，、；;:：（）)\]\[{}+]+$', '', s)
    s = s.strip()
    return s


def _add_unique(result, seen, text):
    """去重追加"""
    key = text.lower()
    if key not in seen:
        seen.add(key)
        result.append(text)


def _extract_entry_from_line(line):
    """
    从单行中提取完整条目（优先保留多词短语/句型/谚语）
    返回条目标题字符串，若无法提取返回 None
    """
    line = line.strip()
    if not line:
        return None
    # 以第一个中文字符为界，隔离出英文部分
    m = re.search(r'[\u4e00-\u9fff]', line)
    has_cn = m is not None
    english_part = line[:m.start()] if m else line
    entry = _clean_entry(english_part)
    if not entry:
        return None
    parts = [p for p in entry.split() if p]
    if has_cn:
        # 明确是单词/短语条目（带中文释义），尽量完整保留
        if len(entry) <= MAX_PHRASE_LEN and ENTRY_PATTERN.match(entry):
            return entry
        return None
    # 纯英文行
    if len(parts) == 1:
        if is_valid_word(entry):
            return entry
        return None
    if is_valid_phrase(entry):
        return entry
    return None


def _extract_tokens_from_text(text, result, seen):
    """从非条目文本中按 token 抽取单词（兜底，处理多词混排的文本）"""
    raw_tokens = re.split(
        r'[\n\r\t,;:。，；：、！？!?"\'\(\)\[\]\{\}<>【】《》·…—\*#+/\\|]+', text)
    for token in raw_tokens:
        token = token.strip()
        if not token:
            continue
        parts = token.split()
        if is_valid_word(token):
            _add_unique(result, seen, token)
            continue
        for part in parts:
            part = re.sub(r'[^a-zA-Z\- ]+$', '', part)
            part = re.sub(r'^[^a-zA-Z\- ]+', '', part)
            if is_valid_word(part):
                _add_unique(result, seen, part)


def extract_words_from_text(text):
    """
    从任意文本中提取单词列表
    优先按行提取完整条目（保留多词短语/固定搭配/句型/谚语），
    无法识别的行再退化为按 token 抽取。
    返回去重后的有序列表（保持出现顺序）
    """
    if not text:
        return []
    lines = text.splitlines()
    seen = set()
    result = []
    for line in lines:
        entry = _extract_entry_from_line(line)
        if entry:
            _add_unique(result, seen, entry)
        else:
            _extract_tokens_from_text(line, result, seen)
    return result


def parse_txt(file_bytes):
    """解析 txt 文件（尝试多种编码）"""
    for encoding in ['utf-8', 'gbk', 'gb2312', 'latin-1']:
        try:
            return file_bytes.decode(encoding)
        except (UnicodeDecodeError, LookupError):
            continue
    return file_bytes.decode('utf-8', errors='ignore')


def parse_docx(file_bytes):
    """解析 Word docx 文件"""
    from docx import Document
    doc = Document(io.BytesIO(file_bytes))
    paragraphs = [p.text for p in doc.paragraphs]
    # 表格中的文字
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                paragraphs.append(cell.text)
    return '\n'.join(paragraphs)


def parse_xlsx(file_bytes):
    """解析 Excel xlsx 文件"""
    from openpyxl import load_workbook
    wb = load_workbook(io.BytesIO(file_bytes), read_only=True, data_only=True)
    texts = []
    for sheet in wb.worksheets:
        for row in sheet.iter_rows(values_only=True):
            for cell in row:
                if cell is not None:
                    texts.append(str(cell))
    wb.close()
    return '\n'.join(texts)


def parse_pdf(file_bytes):
    """解析 PDF 文件"""
    import pdfplumber
    texts = []
    with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
        for page in pdf.pages:
            page_text = page.extract_text() or ''
            texts.append(page_text)
    return '\n'.join(texts)


def parse_document(file_bytes, filename):
    """
    根据文件扩展名解析文档，返回提取的单词列表
    返回: (words_list, raw_text)
    """
    ext = filename.lower().rsplit('.', 1)[-1] if '.' in filename else ''
    if ext == 'txt':
        text = parse_txt(file_bytes)
    elif ext == 'docx':
        text = parse_docx(file_bytes)
    elif ext in ('xlsx', 'xls'):
        text = parse_xlsx(file_bytes)
    elif ext == 'pdf':
        text = parse_pdf(file_bytes)
    else:
        raise ValueError(f'不支持的文件格式：{ext}（支持 txt/docx/xlsx/pdf）')

    words = extract_words_from_text(text)
    return words, text


def parse_document_preview(file_bytes, filename, max_raw=500):
    """
    解析文档并返回预览信息（不导入，仅展示）
    返回: {words, total, filtered, raw_preview}
    """
    words, raw_text = parse_document(file_bytes, filename)
    return {
        'words': words,
        'total': len(words),
        'raw_preview': raw_text[:max_raw] if raw_text else '',
        'raw_length': len(raw_text) if raw_text else 0,
    }